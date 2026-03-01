"""
Phase 2-M9: Word/Excel出力エクスポーター
disclosure-multiagent — m9_document_exporter.py

## 概要
M4の ProposalSet（松竹梅提案セット）を受け取り、
Word（.docx）または Excel（.xlsx）形式でエクスポートする。

## 使い方
    from m9_document_exporter import export_to_word, export_to_excel, export_documents

## 設計書
    00_Requirements_Definition.md §3.2 F-09「差分のワード/Excel出力」

## 依存
    python-docx>=1.0.0  (pip install python-docx)
    openpyxl>=3.0.0     (requirements_poc.txt 参照)
    m4_proposal_agent.py（ProposalSet・Proposal）
"""

from __future__ import annotations

import sys
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Optional

# scripts ディレクトリをパスに追加（同ディレクトリ内の m4 をインポートするため）
_SCRIPTS_DIR = Path(__file__).parent
if str(_SCRIPTS_DIR) not in sys.path:
    sys.path.insert(0, str(_SCRIPTS_DIR))

from m4_proposal_agent import ProposalSet  # noqa: E402

# ------------------------------------------------------------------
# ライブラリ存在チェック（未インストール時はスキップテスト用）
# ------------------------------------------------------------------

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Pt, RGBColor
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False

# ------------------------------------------------------------------
# 定数
# ------------------------------------------------------------------

LEVELS = ("松", "竹", "梅")

# Excel 列ヘッダー（F-09設計書準拠）
EXCEL_HEADERS = [
    "GAP_ID",
    "ギャップ要約",
    "松（テキスト）",
    "竹（テキスト）",
    "梅（テキスト）",
    "法令根拠",
    "警告",
]

# ヘッダー背景色（Excel）
HEADER_FILL_COLOR = "4472C4"   # 青系
HEADER_FONT_COLOR = "FFFFFF"   # 白

# ------------------------------------------------------------------
# データクラス
# ------------------------------------------------------------------


@dataclass
class DocumentExportResult:
    """
    Word/Excelエクスポート結果。
    word_path / excel_path が空文字列の場合は未生成（ライブラリ未インストール等）。
    """
    word_path: str
    excel_path: str
    proposal_count: int
    export_at: str


# ------------------------------------------------------------------
# Word 出力
# ------------------------------------------------------------------


def _add_hyperlink(paragraph, text: str, url: str) -> None:
    """python-docx ヘルパー: 段落内にハイパーリンクを追加する。"""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    r.append(rPr)

    t = OxmlElement("w:t")
    t.text = text
    r.append(t)
    hyperlink.append(r)
    paragraph._p.append(hyperlink)


def export_to_word(
    proposal_sets: list[ProposalSet],
    output_path: str,
    company_name: str = "分析対象企業",
    fiscal_year: int = 0,
) -> str:
    """
    ProposalSet リストを Word (.docx) ファイルに出力する。

    Args:
        proposal_sets: M4 が生成した ProposalSet のリスト
        output_path: 出力先ファイルパス（.docx）
        company_name: 会社名（ヘッダー表示用）
        fiscal_year: 会計年度（0 の場合は表示なし）

    Returns:
        output_path（生成されたファイルパス）

    Raises:
        ImportError: python-docx が未インストールの場合
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx が未インストールです。pip install python-docx を実行してください。")

    doc = Document()

    # ── タイトル ──────────────────────────────────────────────────
    title_para = doc.add_heading("人的資本開示 改善提案レポート", level=1)
    title_para.alignment = 1  # CENTER

    # ── ヘッダー情報 ─────────────────────────────────────────────
    info_para = doc.add_paragraph()
    year_str = f"  会計年度: {fiscal_year}年度" if fiscal_year else ""
    info_para.add_run(
        f"会社名: {company_name}{year_str}  出力日: {datetime.now().strftime('%Y-%m-%d')}"
    ).font.size = Pt(10)

    doc.add_paragraph()  # 空行

    # ── 各 ProposalSet ────────────────────────────────────────────
    for ps in proposal_sets:
        # GAP 見出し
        doc.add_heading(f"{ps.gap_id}  {ps.disclosure_item}", level=2)

        # source_warning
        if ps.source_warning:
            warn_para = doc.add_paragraph()
            run = warn_para.add_run(f"⚠️  {ps.source_warning}")
            run.font.color.rgb = RGBColor(0xFF, 0x99, 0x00)  # オレンジ
            run.font.bold = True

        # 松竹梅 各提案
        for level_label, proposal in [
            ("松", ps.matsu),
            ("竹", ps.take),
            ("梅", ps.ume),
        ]:
            # レベル小見出し
            doc.add_heading(f"【{level_label}】", level=3)

            # 本文（改行対応）
            body_para = doc.add_paragraph()
            body_para.add_run(proposal.text)

            # 法令根拠
            law_para = doc.add_paragraph()
            law_run_label = law_para.add_run("法令根拠: ")
            law_run_label.font.bold = True

            if level_label == "松" and ps.reference_url:
                # 松のみリンク付き
                _add_hyperlink(law_para, ps.reference_law_id, ps.reference_url)
            else:
                law_para.add_run(ps.reference_law_id)

        doc.add_paragraph()  # Gap 間の空行

    doc.save(output_path)
    return output_path


# ------------------------------------------------------------------
# Excel 出力
# ------------------------------------------------------------------


def export_to_excel(
    proposal_sets: list[ProposalSet],
    output_path: str,
    company_name: str = "分析対象企業",
    fiscal_year: int = 0,
) -> str:
    """
    ProposalSet リストを Excel (.xlsx) ファイルに出力する。

    シート「提案一覧」に1行1GAPで出力する。

    Args:
        proposal_sets: M4 が生成した ProposalSet のリスト
        output_path: 出力先ファイルパス（.xlsx）
        company_name: 会社名（シートタイトル用）
        fiscal_year: 会計年度（シートタイトル用）

    Returns:
        output_path（生成されたファイルパス）

    Raises:
        ImportError: openpyxl が未インストールの場合
    """
    if not OPENPYXL_AVAILABLE:
        raise ImportError("openpyxl が未インストールです。pip install openpyxl を実行してください。")

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "提案一覧"

    # ── シートタイトル行 ───────────────────────────────────────────
    year_str = f"  会計年度: {fiscal_year}年度" if fiscal_year else ""
    ws.merge_cells("A1:G1")
    title_cell = ws["A1"]
    title_cell.value = f"人的資本開示 改善提案レポート  会社名: {company_name}{year_str}  出力日: {datetime.now().strftime('%Y-%m-%d')}"
    title_cell.font = Font(bold=True, size=12)
    title_cell.alignment = Alignment(horizontal="center")
    ws.row_dimensions[1].height = 20

    # ── ヘッダー行 ────────────────────────────────────────────────
    header_font = Font(bold=True, color=HEADER_FONT_COLOR, size=10)
    header_fill = PatternFill("solid", fgColor=HEADER_FILL_COLOR)
    header_align = Alignment(horizontal="center", vertical="center", wrap_text=True)

    for col_idx, header in enumerate(EXCEL_HEADERS, start=1):
        cell = ws.cell(row=2, column=col_idx, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_align

    ws.row_dimensions[2].height = 24

    # ── データ行 ─────────────────────────────────────────────────
    data_align = Alignment(vertical="top", wrap_text=True)

    for row_idx, ps in enumerate(proposal_sets, start=3):
        warning_text = f"⚠️ {ps.source_warning}" if ps.source_warning else ""

        row_data = [
            ps.gap_id,
            ps.disclosure_item,
            ps.matsu.text,
            ps.take.text,
            ps.ume.text,
            ps.reference_law_id,
            warning_text,
        ]

        for col_idx, value in enumerate(row_data, start=1):
            cell = ws.cell(row=row_idx, column=col_idx, value=value)
            cell.alignment = data_align

    # ── 列幅設定 ─────────────────────────────────────────────────
    col_widths = [12, 28, 45, 35, 25, 18, 22]
    for col_idx, width in enumerate(col_widths, start=1):
        ws.column_dimensions[openpyxl.utils.get_column_letter(col_idx)].width = width

    wb.save(output_path)
    return output_path


# ------------------------------------------------------------------
# 統合エクスポート関数
# ------------------------------------------------------------------


def export_documents(
    proposal_sets: list[ProposalSet],
    word_path: str,
    excel_path: str,
    company_name: str = "分析対象企業",
    fiscal_year: int = 0,
) -> DocumentExportResult:
    """
    Word と Excel を一括エクスポートし DocumentExportResult を返す。

    どちらかのライブラリが未インストールの場合は、
    そのパスを空文字列にして部分成功とする（ImportError を上げない）。
    """
    actual_word_path = ""
    actual_excel_path = ""

    if DOCX_AVAILABLE:
        export_to_word(proposal_sets, word_path, company_name, fiscal_year)
        actual_word_path = word_path

    if OPENPYXL_AVAILABLE:
        export_to_excel(proposal_sets, excel_path, company_name, fiscal_year)
        actual_excel_path = excel_path

    return DocumentExportResult(
        word_path=actual_word_path,
        excel_path=actual_excel_path,
        proposal_count=len(proposal_sets),
        export_at=datetime.now().strftime("%Y-%m-%dT%H:%M:%S"),
    )
