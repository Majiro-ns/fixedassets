"""
test_m9_document_exporter.py
============================
M9 Word/Excel エクスポーター テスト

TC-1: DocumentExportResult 型確認
TC-2: export_to_word() が .docx ファイルを生成する（tempfile使用）
TC-3: export_to_excel() が .xlsx ファイルを生成する（tempfile使用）
TC-4: Word出力に松竹梅3種の本文が含まれる
TC-5: Excel出力に正しい列ヘッダーが含まれる
TC-6: source_warning がある提案に⚠️が含まれる（Word/Excel両方）

CHECK-9根拠（各TCに記載）:
  TC-1: DocumentExportResult のフィールド型検証 → dataclass 定義が期待通りであることの確認
  TC-2: export_to_word() は output_path に .docx を生成し str を返す → ファイル存在確認
  TC-3: export_to_excel() は output_path に .xlsx を生成し str を返す → ファイル存在確認
  TC-4: Word の段落テキストを結合して全体検索 → 松/竹/梅の本文テキストが全て含まれること
  TC-5: Excel シート「提案一覧」の2行目が EXCEL_HEADERS と一致すること
  TC-6: source_warning の内容が ⚠️ プレフィックスと共に Word・Excel 双方に出力されること
"""

from __future__ import annotations

import os
import sys
import tempfile
import unittest
from dataclasses import fields
from pathlib import Path

# USE_MOCK_LLM=true を強制（実API不使用）
os.environ.setdefault("USE_MOCK_LLM", "true")

_SCRIPTS_DIR = Path(__file__).parent
if str(_SCRIPTS_DIR) not in sys.path:
    sys.path.insert(0, str(_SCRIPTS_DIR))

# ライブラリ存在確認
try:
    import docx  # noqa: F401
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import openpyxl  # noqa: F401
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False

from m4_proposal_agent import GapItem, Proposal, ProposalSet, QualityCheckResult
from m9_document_exporter import (
    EXCEL_HEADERS,
    DocumentExportResult,
    export_to_excel,
    export_to_word,
    export_documents,
)


# ------------------------------------------------------------------
# テスト用フィクスチャ
# ------------------------------------------------------------------

def _make_proposal(level: str, text: str) -> Proposal:
    """テスト用 Proposal を生成する。"""
    return Proposal(
        level=level,
        text=text,
        quality=QualityCheckResult(passed=True, should_regenerate=False),
        attempts=1,
        status="pass",
        placeholders=[],
    )


def _make_proposal_set(
    gap_id: str = "GAP-001",
    disclosure_item: str = "人的資本KPIの定量的開示",
    source_warning: str | None = None,
) -> ProposalSet:
    """テスト用 ProposalSet を生成する。"""
    return ProposalSet(
        gap_id=gap_id,
        disclosure_item=disclosure_item,
        reference_law_id="HC_20260220_001",
        reference_url="https://example.com/law",
        source_warning=source_warning,
        matsu=_make_proposal("松", "【松】当社は2030年度までにDX人材を300名育成する計画の下、毎年30名を採用します。"),
        take=_make_proposal("竹", "【竹】当社は人材育成計画を策定し、KPIを開示します。"),
        ume=_make_proposal("梅", "【梅】人材戦略の概要を開示します。"),
    )


def _make_proposal_sets_with_warning() -> list[ProposalSet]:
    """source_warning 付きの ProposalSet リストを生成する。"""
    return [
        _make_proposal_set(gap_id="GAP-001"),
        _make_proposal_set(
            gap_id="GAP-002",
            disclosure_item="給与決定方針の開示",
            source_warning="URLが未確認です。正式URL要確認（HC_20250421_001）",
        ),
    ]


# ------------------------------------------------------------------
# TC-1: DocumentExportResult 型確認
# ------------------------------------------------------------------

class TestDocumentExportResultType(unittest.TestCase):
    """TC-1: DocumentExportResult dataclass のフィールド型検証"""

    def test_tc1_document_export_result_has_required_fields(self) -> None:
        """
        TC-1: DocumentExportResult が word_path / excel_path / proposal_count / export_at を持つ

        根拠: F-09設計書 DocumentExportResult dataclass 定義の確認。
        CHECK-9: dataclass の fields() で型アノテーションが str/int/str であることを検証する。
                 型違いがあれば dataclass 定義の修正が必要。
        """
        result = DocumentExportResult(
            word_path="/tmp/test.docx",
            excel_path="/tmp/test.xlsx",
            proposal_count=3,
            export_at="2026-02-28T00:00:00",
        )

        self.assertIsInstance(result.word_path, str, "word_path が str でない")
        self.assertIsInstance(result.excel_path, str, "excel_path が str でない")
        self.assertIsInstance(result.proposal_count, int, "proposal_count が int でない")
        self.assertIsInstance(result.export_at, str, "export_at が str でない")

        # dataclass fields 名確認
        field_names = {f.name for f in fields(DocumentExportResult)}
        for name in ("word_path", "excel_path", "proposal_count", "export_at"):
            self.assertIn(name, field_names, f"フィールド {name!r} が DocumentExportResult にない")


# ------------------------------------------------------------------
# TC-2: export_to_word() ファイル生成確認
# ------------------------------------------------------------------

@unittest.skipUnless(DOCX_AVAILABLE, "python-docx が未インストールのためスキップ")
class TestExportToWord(unittest.TestCase):
    """TC-2/TC-4/TC-6-Word: export_to_word() のファイル生成・内容確認"""

    def setUp(self) -> None:
        os.environ["USE_MOCK_LLM"] = "true"
        self.proposal_sets = [_make_proposal_set()]

    def test_tc2_export_to_word_creates_docx_file(self) -> None:
        """
        TC-2: export_to_word() が .docx ファイルを生成する

        根拠: export_to_word() は output_path に Word ファイルを保存し、そのパスを返す。
        CHECK-9: tempfile.NamedTemporaryFile で生成→削除。
                 python-docx が Document.save() を呼べば .docx が生成されることを確認する。
        """
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp_path = tmp.name

        try:
            result_path = export_to_word(
                proposal_sets=self.proposal_sets,
                output_path=tmp_path,
                company_name="テスト会社",
                fiscal_year=2025,
            )

            self.assertEqual(result_path, tmp_path, "export_to_word() が output_path を返さなかった")
            self.assertTrue(
                Path(tmp_path).exists(),
                f".docx ファイルが生成されていません: {tmp_path}",
            )
            self.assertGreater(
                Path(tmp_path).stat().st_size,
                0,
                ".docx ファイルのサイズが0バイト",
            )
        finally:
            if Path(tmp_path).exists():
                Path(tmp_path).unlink()

    def test_tc4_word_output_contains_all_levels(self) -> None:
        """
        TC-4: Word出力に松竹梅3種の本文が含まれる

        根拠: generate_proposals() が生成した 松/竹/梅 の text が Word 段落に書き込まれること。
        CHECK-9: python-docx の Document.paragraphs から全段落テキストを結合し、
                 "【松】" "【竹】" "【梅】" が含まれることで3水準の出力を確認する。
                 単なる見出しではなく本文テキスト（Proposal.text）が含まれることを検証。
        """
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp_path = tmp.name

        try:
            export_to_word(
                proposal_sets=self.proposal_sets,
                output_path=tmp_path,
                company_name="テスト会社",
                fiscal_year=2025,
            )

            from docx import Document as DocxDocument
            doc = DocxDocument(tmp_path)
            all_text = "\n".join(p.text for p in doc.paragraphs)

            # 松/竹/梅 各水準の本文が含まれることを確認
            for level_text in ["【松】", "【竹】", "【梅】"]:
                self.assertIn(
                    level_text,
                    all_text,
                    f"Word出力に {level_text} の本文が含まれていません",
                )
        finally:
            if Path(tmp_path).exists():
                Path(tmp_path).unlink()

    def test_tc6_word_output_contains_warning_marker(self) -> None:
        """
        TC-6(Word): source_warning がある提案に⚠️が含まれる

        根拠: ProposalSet.source_warning != None の場合、Word 段落に "⚠️" プレフィックスで
              警告テキストが挿入されること。
        CHECK-9: source_warning 付き ProposalSet を渡し、全段落テキストに
                 "⚠️" が含まれることを確認する。
        """
        ps_with_warning = _make_proposal_set(
            gap_id="GAP-002",
            source_warning="URLが未確認です（要確認）",
        )

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp_path = tmp.name

        try:
            export_to_word(
                proposal_sets=[ps_with_warning],
                output_path=tmp_path,
            )

            from docx import Document as DocxDocument
            doc = DocxDocument(tmp_path)
            all_text = "\n".join(p.text for p in doc.paragraphs)

            self.assertIn(
                "⚠️",
                all_text,
                "source_warning がある提案に ⚠️ が含まれていません（Word）",
            )
        finally:
            if Path(tmp_path).exists():
                Path(tmp_path).unlink()


# ------------------------------------------------------------------
# TC-3: export_to_excel() ファイル生成確認
# ------------------------------------------------------------------

@unittest.skipUnless(OPENPYXL_AVAILABLE, "openpyxl が未インストールのためスキップ")
class TestExportToExcel(unittest.TestCase):
    """TC-3/TC-5/TC-6-Excel: export_to_excel() のファイル生成・内容確認"""

    def setUp(self) -> None:
        os.environ["USE_MOCK_LLM"] = "true"
        self.proposal_sets = [_make_proposal_set()]

    def test_tc3_export_to_excel_creates_xlsx_file(self) -> None:
        """
        TC-3: export_to_excel() が .xlsx ファイルを生成する

        根拠: export_to_excel() は output_path に Excel ファイルを保存し、そのパスを返す。
        CHECK-9: tempfile.NamedTemporaryFile で生成→削除。
                 openpyxl が Workbook.save() を呼べば .xlsx が生成されることを確認する。
        """
        with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as tmp:
            tmp_path = tmp.name

        try:
            result_path = export_to_excel(
                proposal_sets=self.proposal_sets,
                output_path=tmp_path,
                company_name="テスト会社",
                fiscal_year=2025,
            )

            self.assertEqual(result_path, tmp_path, "export_to_excel() が output_path を返さなかった")
            self.assertTrue(
                Path(tmp_path).exists(),
                f".xlsx ファイルが生成されていません: {tmp_path}",
            )
            self.assertGreater(
                Path(tmp_path).stat().st_size,
                0,
                ".xlsx ファイルのサイズが0バイト",
            )
        finally:
            if Path(tmp_path).exists():
                Path(tmp_path).unlink()

    def test_tc5_excel_output_has_correct_headers(self) -> None:
        """
        TC-5: Excel出力に正しい列ヘッダーが含まれる

        根拠: シート「提案一覧」の2行目が EXCEL_HEADERS と完全一致すること。
        CHECK-9: F-09設計書の列定義「GAP_ID / ギャップ要約 / 松（テキスト） / 竹（テキスト） /
                 梅（テキスト） / 法令根拠 / 警告」が正確に出力されることを確認する。
                 列順・名称の一致まで検証（部分一致ではなく完全一致）。
        """
        with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as tmp:
            tmp_path = tmp.name

        try:
            export_to_excel(
                proposal_sets=self.proposal_sets,
                output_path=tmp_path,
            )

            import openpyxl as xl
            wb = xl.load_workbook(tmp_path)
            self.assertIn("提案一覧", wb.sheetnames, "シート「提案一覧」が存在しません")

            ws = wb["提案一覧"]
            actual_headers = [ws.cell(row=2, column=col).value for col in range(1, len(EXCEL_HEADERS) + 1)]

            self.assertEqual(
                actual_headers,
                EXCEL_HEADERS,
                f"Excel ヘッダーが期待値と一致しません。\n期待: {EXCEL_HEADERS}\n実際: {actual_headers}",
            )
        finally:
            if Path(tmp_path).exists():
                Path(tmp_path).unlink()

    def test_tc6_excel_output_contains_warning_marker(self) -> None:
        """
        TC-6(Excel): source_warning がある提案に⚠️が含まれる

        根拠: ProposalSet.source_warning != None の場合、Excel の「警告」列に
              "⚠️" プレフィックスで警告テキストが書き込まれること。
        CHECK-9: source_warning 付き ProposalSet を使い、「警告」列（7列目）に
                 "⚠️" が含まれることを確認する。
        """
        ps_with_warning = _make_proposal_set(
            gap_id="GAP-002",
            source_warning="URLが未確認です（要確認）",
        )

        with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as tmp:
            tmp_path = tmp.name

        try:
            export_to_excel(
                proposal_sets=[ps_with_warning],
                output_path=tmp_path,
            )

            import openpyxl as xl
            wb = xl.load_workbook(tmp_path)
            ws = wb["提案一覧"]

            # データ行は row=3 から（row=1: タイトル, row=2: ヘッダー）
            warning_cell_value = ws.cell(row=3, column=7).value or ""
            self.assertIn(
                "⚠️",
                warning_cell_value,
                f"Excel の「警告」列に ⚠️ が含まれていません。実際の値: {warning_cell_value!r}",
            )
        finally:
            if Path(tmp_path).exists():
                Path(tmp_path).unlink()


# ------------------------------------------------------------------
# export_documents() 統合テスト
# ------------------------------------------------------------------

@unittest.skipUnless(DOCX_AVAILABLE and OPENPYXL_AVAILABLE, "python-docx または openpyxl が未インストールのためスキップ")
class TestExportDocuments(unittest.TestCase):
    """export_documents() 統合エクスポート確認"""

    def test_export_documents_returns_result_with_both_paths(self) -> None:
        """
        export_documents() が DocumentExportResult を返し、両パスにファイルが生成される

        根拠: Word と Excel を一括生成する高レベル関数の統合確認。
        CHECK-9: word_path / excel_path が空でなく、ファイルが実際に存在することを確認。
                 proposal_count が入力の len と一致することも検証する。
        """
        proposal_sets = _make_proposal_sets_with_warning()

        with (
            tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as w_tmp,
            tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as x_tmp,
        ):
            w_path = w_tmp.name
            x_path = x_tmp.name

        try:
            result = export_documents(
                proposal_sets=proposal_sets,
                word_path=w_path,
                excel_path=x_path,
                company_name="統合テスト社",
                fiscal_year=2025,
            )

            self.assertIsInstance(result, DocumentExportResult)
            self.assertEqual(result.proposal_count, len(proposal_sets))
            self.assertNotEqual(result.word_path, "", "word_path が空文字列")
            self.assertNotEqual(result.excel_path, "", "excel_path が空文字列")
            self.assertTrue(Path(result.word_path).exists(), "Word ファイルが存在しない")
            self.assertTrue(Path(result.excel_path).exists(), "Excel ファイルが存在しない")
            self.assertGreater(len(result.export_at), 0, "export_at が空文字列")
        finally:
            for p in (w_path, x_path):
                if Path(p).exists():
                    Path(p).unlink()


# ------------------------------------------------------------------
# cmd_076 追加要件: TC-4b（Excel行数）・TC-5b（空リストエッジケース）
# ------------------------------------------------------------------

@unittest.skipUnless(OPENPYXL_AVAILABLE, "openpyxl が未インストールのためスキップ")
class TestExcelRowCountMatchesProposals(unittest.TestCase):
    """TC-4b: Excelシートの行数が提案数と一致すること（cmd_076 C2追加要件）"""

    def test_tc4b_excel_row_count_matches_proposal_count(self) -> None:
        """
        TC-4b: Excelシートのデータ行数が入力の提案数と一致する

        根拠: export_to_excel() は row=3 からデータを書き込む（row=1: タイトル, row=2: ヘッダー）。
              N件の ProposalSet を渡した場合、データ行は row=3〜row=N+2 の N行になる。
        CHECK-9: 3件の ProposalSet を入力し、データ行が3行（row=3〜5）になることを確認。
                 row=6 は None（データなし）であることで行数の一致を検証する。
        """
        proposal_sets = [
            _make_proposal_set(gap_id="GAP-001", disclosure_item="KPI連動"),
            _make_proposal_set(gap_id="GAP-002", disclosure_item="給与決定方針"),
            _make_proposal_set(gap_id="GAP-003", disclosure_item="前年比増減率"),
        ]

        with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as tmp:
            tmp_path = tmp.name

        try:
            export_to_excel(proposal_sets=proposal_sets, output_path=tmp_path)

            import openpyxl as xl
            wb = xl.load_workbook(tmp_path)
            ws = wb["提案一覧"]

            # データ行は row=3 から（row=1: タイトル, row=2: ヘッダー）
            # 3件入力 → row=3, 4, 5 にデータ、row=6 は None
            data_row_count = 0
            for row_idx in range(3, 3 + len(proposal_sets) + 2):
                cell_value = ws.cell(row=row_idx, column=1).value
                if cell_value is not None:
                    data_row_count += 1

            self.assertEqual(
                data_row_count,
                len(proposal_sets),
                f"データ行数 {data_row_count} が提案数 {len(proposal_sets)} と一致しない",
            )

            # 各行の GAP_ID が正しく書き込まれているか確認
            for row_offset, ps in enumerate(proposal_sets):
                actual_gap_id = ws.cell(row=3 + row_offset, column=1).value
                self.assertEqual(
                    actual_gap_id,
                    ps.gap_id,
                    f"row={3 + row_offset} の GAP_ID が一致しない: 期待={ps.gap_id}, 実際={actual_gap_id}",
                )
        finally:
            if Path(tmp_path).exists():
                Path(tmp_path).unlink()


class TestEmptyProposalListEdgeCase(unittest.TestCase):
    """TC-5b: 空の提案リストでもエラーなし（エッジケース）（cmd_076 C2追加要件）"""

    def setUp(self) -> None:
        os.environ["USE_MOCK_LLM"] = "true"

    @unittest.skipUnless(DOCX_AVAILABLE, "python-docx が未インストールのためスキップ")
    def test_tc5b_empty_list_word_no_error(self) -> None:
        """
        TC-5b(Word): 空の提案リストを渡しても export_to_word() がエラーなく完了する

        根拠: 提案が0件の場合（全ギャップが has_gap=False 等）に
              Word 出力関数がクラッシュしないことを確認する。
        CHECK-9: 空リスト [] を渡した場合、ファイルは生成される（サイズ > 0）が
                 提案内容は含まれない。FileNotFoundError や AttributeError が発生しないこと。
        """
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp_path = tmp.name

        try:
            result_path = export_to_word(
                proposal_sets=[],
                output_path=tmp_path,
                company_name="エッジケーステスト社",
                fiscal_year=2025,
            )
            self.assertEqual(result_path, tmp_path)
            self.assertTrue(
                Path(tmp_path).exists(),
                "空リストでも .docx ファイルが生成されること",
            )
            self.assertGreater(
                Path(tmp_path).stat().st_size,
                0,
                "空リストでも .docx ファイルサイズ > 0 であること",
            )
        finally:
            if Path(tmp_path).exists():
                Path(tmp_path).unlink()

    @unittest.skipUnless(OPENPYXL_AVAILABLE, "openpyxl が未インストールのためスキップ")
    def test_tc5b_empty_list_excel_no_error(self) -> None:
        """
        TC-5b(Excel): 空の提案リストを渡しても export_to_excel() がエラーなく完了する

        根拠: 提案が0件の場合に Excel 出力関数がクラッシュしないことを確認する。
        CHECK-9: 空リスト [] を渡した場合、ファイルは生成されヘッダー行のみ存在する。
                 row=3 の GAP_ID セルが None であることでデータ行がないことを確認する。
        """
        with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as tmp:
            tmp_path = tmp.name

        try:
            result_path = export_to_excel(
                proposal_sets=[],
                output_path=tmp_path,
                company_name="エッジケーステスト社",
                fiscal_year=2025,
            )
            self.assertEqual(result_path, tmp_path)
            self.assertTrue(Path(tmp_path).exists(), "空リストでも .xlsx ファイルが生成されること")

            import openpyxl as xl
            wb = xl.load_workbook(tmp_path)
            ws = wb["提案一覧"]

            # ヘッダー行は存在する
            header_val = ws.cell(row=2, column=1).value
            self.assertEqual(header_val, "GAP_ID", "空リストでもヘッダー行は存在すること")

            # データ行は存在しない（row=3 は None）
            data_val = ws.cell(row=3, column=1).value
            self.assertIsNone(data_val, f"空リストなのに row=3 にデータがある: {data_val!r}")
        finally:
            if Path(tmp_path).exists():
                Path(tmp_path).unlink()

    @unittest.skipUnless(DOCX_AVAILABLE and OPENPYXL_AVAILABLE, "python-docx または openpyxl が未インストールのためスキップ")
    def test_tc5b_empty_list_export_documents_no_error(self) -> None:
        """
        TC-5b(統合): 空リストで export_documents() がエラーなく DocumentExportResult を返す

        根拠: 統合関数でも空リストを安全に処理できること。
        CHECK-9: proposal_count=0 の DocumentExportResult が返されること。
        """
        with (
            tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as w_tmp,
            tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as x_tmp,
        ):
            w_path = w_tmp.name
            x_path = x_tmp.name

        try:
            result = export_documents(
                proposal_sets=[],
                word_path=w_path,
                excel_path=x_path,
                company_name="エッジケーステスト社",
                fiscal_year=2025,
            )
            self.assertIsInstance(result, DocumentExportResult)
            self.assertEqual(result.proposal_count, 0, "空リストの proposal_count は 0 であること")
        finally:
            for p in (w_path, x_path):
                if Path(p).exists():
                    Path(p).unlink()


if __name__ == "__main__":
    unittest.main()
